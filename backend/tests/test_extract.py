"""Extracción de PDF/DOCX robusta: leer el documento COMPLETO (no 'cachitos'),
tolerar una página que falla, y avisar cuando un PDF es escaneado.

Requiere fpdf2 y python-docx (los del propio proyecto). Correr:
    python tests/test_extract.py
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))


def _make_pdf(n_pages: int) -> bytes:
    from fpdf import FPDF
    import io
    p = FPDF()
    for pag in range(1, n_pages + 1):
        p.add_page()
        p.set_font("helvetica", size=12)
        p.multi_cell(0, 6, f"Pagina numero {pag}. " +
                     " ".join(f"palabra{pag}n{i}" for i in range(35)))
    buf = io.BytesIO()
    p.output(buf)
    return buf.getvalue()


def test_pdf_full_extraction():
    from app import extract
    data = _make_pdf(5)
    info = extract.extract_detailed("multi.pdf", data)
    assert info["pages_total"] == 5
    assert info["pages_with_text"] == 5
    assert info["partial"] is False
    for i in range(1, 6):
        assert f"Pagina numero {i}" in info["text"], f"Falta la página {i}"


def test_pdf_bad_page_does_not_lose_rest():
    """Si una página falla, se recuperan las demás (antes se perdía todo)."""
    from unittest.mock import patch
    import pypdf
    from app import extract
    data = _make_pdf(5)
    orig = pypdf._page.PageObject.extract_text

    def flaky(self, *a, **k):
        if "Pagina numero 3" in (orig(self) or ""):
            raise RuntimeError("página corrupta")
        return orig(self)

    with patch.object(pypdf._page.PageObject, "extract_text", flaky):
        info = extract.extract_detailed("multi.pdf", data)
    present = [i for i in range(1, 6) if f"Pagina numero {i}" in info["text"]]
    assert present == [1, 2, 4, 5], f"Debería recuperar 4 páginas, tiene {present}"
    assert info["pages_with_text"] == 4 and info["partial"] is True


def test_pdf_scanned_is_reported():
    """Un PDF sin texto (solo imágenes) se reporta como escaneado, no rompe."""
    from unittest.mock import patch
    import pypdf
    from app import extract
    data = _make_pdf(3)
    with patch.object(pypdf._page.PageObject, "extract_text",
                      lambda self, *a, **k: ""):
        info = extract.extract_detailed("scan.pdf", data)
    assert info["text"] == ""
    assert info["scanned"] is True
    assert "escane" in (info["note"] or "").lower()


def test_docx_includes_tables():
    from docx import Document
    import io
    from app import extract
    d = Document()
    d.add_paragraph("Primer parrafo del documento de prueba.")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "CeldaUno"
    t.rows[0].cells[1].text = "CeldaDos"
    buf = io.BytesIO()
    d.save(buf)
    info = extract.extract_detailed("doc.docx", buf.getvalue())
    assert "Primer parrafo" in info["text"]
    assert "CeldaUno" in info["text"] and "CeldaDos" in info["text"]


if __name__ == "__main__":
    import importlib.util
    for mod in ("fpdf", "pypdf", "docx"):
        if importlib.util.find_spec(mod) is None:
            print(f"SKIP: falta {mod} en este entorno.")
            sys.exit(0)
    test_pdf_full_extraction()
    print("[OK] PDF de 5 páginas se extrae ENTERO")
    test_pdf_bad_page_does_not_lose_rest()
    print("[OK] una página que falla no pierde el resto del documento")
    test_pdf_scanned_is_reported()
    print("[OK] PDF escaneado se reporta claramente sin romper")
    test_docx_includes_tables()
    print("[OK] DOCX incluye también el texto de las tablas")
    print("\nTODAS LAS PRUEBAS DE EXTRACCION PASARON [OK]")
