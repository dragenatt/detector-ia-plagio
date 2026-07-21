"""Extracción de texto desde archivos subidos (.txt, .md, .pdf, .docx).

Las dependencias (pypdf, python-docx) se importan de forma perezosa para que
el resto del backend funcione aunque no estén instaladas.

`extract_detailed` devuelve el texto MÁS un diagnóstico (páginas leídas de un
total, palabras, si parece escaneado o incompleto), para que la interfaz avise
cuando la extracción no quedó completa. `extract_text` se mantiene para quien
solo necesita el texto (corpus de referencia, etc.).
"""
from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger("veraz.extract")

# Debajo de esta densidad, sospechamos extracción parcial o PDF escaneado.
MIN_WORDS_PER_PAGE = 20


def extract_detailed(filename: str, data: bytes) -> dict:
    """Extrae texto y devuelve un diagnóstico.

    Devuelve dict con:
      text, words, pages_total, pages_with_text, scanned (bool|None),
      partial (bool), note (str|None)
    """
    ext = Path(filename).suffix.lower()

    if ext in (".txt", ".md"):
        text = _decode_text(data)
        return _pack(text)

    if ext == ".pdf":
        text, total, with_text = _extract_pdf(data)
        info = _pack(text, pages_total=total, pages_with_text=with_text)
        # Diagnóstico específico de PDF.
        if total and with_text == 0:
            info["scanned"] = True
            info["note"] = ("Este PDF parece ESCANEADO (imágenes, sin texto "
                            "seleccionable): no se pudo extraer texto. "
                            "Necesitarías un PDF con texto o pasarlo por OCR.")
        elif total:
            info["scanned"] = False
            wpp = info["words"] / max(total, 1)
            if with_text < total or wpp < MIN_WORDS_PER_PAGE:
                info["partial"] = True
                info["note"] = (
                    f"Se extrajo texto de {with_text} de {total} página(s) "
                    f"({info['words']} palabras). La extracción pudo quedar "
                    "incompleta: revisa que el PDF tenga texto seleccionable.")
            else:
                info["note"] = (f"Se leyeron {with_text} de {total} páginas · "
                                f"{info['words']} palabras.")
        return info

    if ext == ".docx":
        text = _extract_docx(data)
        return _pack(text)

    raise ValueError(f"Formato no soportado: {ext}")


def extract_text(filename: str, data: bytes) -> str:
    """Solo el texto (compatibilidad con usos que no necesitan diagnóstico)."""
    return extract_detailed(filename, data)["text"]


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #

def _pack(text: str, pages_total: int | None = None,
          pages_with_text: int | None = None) -> dict:
    text = (text or "").strip()
    return {
        "text": text,
        "words": len(text.split()),
        "pages_total": pages_total,
        "pages_with_text": pages_with_text,
        "scanned": None,
        "partial": False,
        "note": None,
    }


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_pdf(data: bytes) -> tuple[str, int, int]:
    """Extrae el texto de TODAS las páginas de forma tolerante.

    Si una página falla, se salta y se continúa con las demás (antes, una sola
    página problemática abortaba el documento entero -> se leían "cachitos").
    Devuelve (texto, páginas_totales, páginas_con_texto).
    """
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("Falta la librería 'pypdf' (pip install pypdf).") from e

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:
        raise ValueError(f"El PDF está dañado o no se puede abrir: {e}") from e

    pages = reader.pages
    total = len(pages)
    parts: list[str] = []
    with_text = 0
    for i, page in enumerate(pages):
        text = ""
        try:
            text = page.extract_text() or ""
        except Exception as e:
            logger.warning("PDF: no se pudo leer la página %d/%d: %s",
                           i + 1, total, e)
            # Segundo intento con modo 'layout' (a veces recupera columnas).
            try:
                text = page.extract_text(extraction_mode="layout") or ""
            except Exception:
                text = ""
        text = text.strip()
        if text:
            with_text += 1
            parts.append(text)
    return "\n\n".join(parts).strip(), total, with_text


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("Falta la librería 'python-docx' (pip install python-docx).") from e
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as e:
        raise ValueError(f"El DOCX está dañado o no se puede abrir: {e}") from e
    # Párrafos + texto de tablas (que python-docx no incluye en .paragraphs).
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()
